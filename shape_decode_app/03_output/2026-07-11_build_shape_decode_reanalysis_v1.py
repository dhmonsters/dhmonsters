# Shape Decode 3자 재분석 Markdown을 검수 가능한 Word 보고서로 만든다.
from __future__ import annotations

import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


APP = Path(__file__).resolve().parents[1]
OUT = APP / "03_output"
SOURCE = OUT / "2026-07-11_shape_decode_three_way_reanalysis_v1.md"
DOCX = OUT / "2026-07-11_shape_decode_three_way_reanalysis_v4.docx"
VIDEO_FRAME = OUT / "2026-07-10_ffmpeg_frames_v1" / "frame_012.png"
TARGET_TIMELINE = OUT / "2026-07-10_video_white_to_transparent_target_zoom_v1.png"
MOTION_TRACE = OUT / "2026-07-11_reference_video_motion_trace_v2.png"
FULL_COMPARE = OUT / "2026-07-11_video_vs_current_same_frame_v3.png"
KEY_COMPARE = OUT / "2026-07-11_video_vs_current_key_frames_v1.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GRID = "CBD5E1"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_inline(paragraph, text: str, size: float = 11) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", size - 0.5, INK)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF2F6")
            run._element.get_or_add_rPr().append(shd)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def next_numbering_id(numbering) -> tuple[int, int]:
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    return max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1


def create_numbering(doc: Document, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = next_numbering_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_paragraph(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    add_inline(paragraph, text)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED)


def add_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def build_key_comparison() -> None:
    with Image.open(FULL_COMPARE) as image:
        row_height = image.height // 8
        star = image.crop((0, 0, image.width, row_height))
        circle = image.crop((0, row_height * 5, image.width, row_height * 6))
        output = Image.new("RGB", (image.width, row_height * 2), "white")
        output.paste(star, (0, 0))
        output.paste(circle, (0, row_height))
        output.save(KEY_COMPARE)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("SHAPE DECODE · 3자 재분석"), size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("2026-07-11  ·  "), size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_masthead(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(kicker.add_run("TECHNICAL FINDINGS"), size=9.5, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("Shape Decode 3자 재분석"), size=24, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("원본 영상 · 현재 Python · Claude WebGL 비교와 영상 복제 기준"), size=12.5, color=MUTED)
    for label, value in (
        ("기준", "Downloads의 녹화_2026_07_09_00_28_18_662_trim.mp4"),
        ("목적", "이전 요구사항 혼선을 분리하고 다음 구현의 단일 기준을 확정"),
        ("권고", "영상 시간표를 복원하고 도형 종류만 별·원·네모·세모로 제한"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}  "), size=10, color=INK, bold=True)
        set_run_font(paragraph.add_run(value), size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    column_count = len(rows[0])
    if column_count == 4:
        widths = [980, 2480, 2740, 3160]
        font_size = 8.8
    elif column_count == 2:
        widths = [2200, 7160]
        font_size = 10
    else:
        base = TABLE_WIDTH // column_count
        widths = [base] * column_count
        widths[-1] += TABLE_WIDTH - sum(widths)
        font_size = 9.2
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_inline(paragraph, value, font_size)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_markdown(doc: Document, lines: list[str]) -> None:
    index = 0
    ordered_num_id = None
    bullet_num_id = None
    inserted_first_visual = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            ordered_num_id = None
            bullet_num_id = None
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("## "):
            title = stripped[3:]
            if title == "세 구현 비교":
                doc.add_page_break()
            doc.add_heading(title, level=1)
            if title == "영상에서 직접 확인한 사실":
                add_picture(doc, TARGET_TIMELINE, 6.15, "영상 프레임에서 확인한 흰 타겟의 이동과 투명화")
                add_picture(doc, MOTION_TRACE, 6.15, "OpenCV 특징점으로 추적한 배경의 시계방향 평행 공전 궤적")
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = []
            for table_line in table_lines:
                cells = [cell.strip() for cell in table_line.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                rows.append(cells)
            add_markdown_table(doc, rows)
            if rows and rows[0][:2] == ["항목", "원본 영상"]:
                add_picture(doc, KEY_COMPARE, 6.15, "같은 Frame 값에서 본 영상과 현재 Python 출력. 위는 별, 아래는 원이다")
            continue
        unordered = re.match(r"^-\s+(.*)$", stripped)
        if unordered:
            if bullet_num_id is None:
                bullet_num_id = create_numbering(doc, False)
            add_list_paragraph(doc, unordered.group(1), bullet_num_id)
            index += 1
            continue
        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered:
            if ordered_num_id is None:
                ordered_num_id = create_numbering(doc, True)
            add_list_paragraph(doc, ordered.group(1), ordered_num_id)
            index += 1
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        add_inline(paragraph, stripped)
        if not inserted_first_visual and stripped.startswith("우리가 만들려는 프로그램은"):
            add_picture(doc, VIDEO_FRAME, 6.15, "원본 영상의 기본 화면. 왼쪽 조작부와 394x264 렌더 캔버스")
            inserted_first_visual = True
        index += 1


def audit_document(doc: Document) -> None:
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert all(round(value.inches, 2) == 1.0 for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin))
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        assert tbl_pr.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"
        assert tbl_pr.find(qn("w:tblW")).get(qn("w:type")) == "dxa"
        assert tbl_pr.find(qn("w:tblInd")).get(qn("w:w")) == str(TABLE_INDENT)


def main() -> None:
    build_key_comparison()
    doc = Document()
    configure_document(doc)
    add_masthead(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()[3:]
    parse_markdown(doc, lines)
    audit_document(doc)
    doc.core_properties.title = "Shape Decode 영상·Python·Claude 3자 재분석"
    doc.core_properties.subject = "영상 복제 기준과 구현 차이 분석"
    doc.core_properties.author = "Codex"
    doc.save(DOCX)
    print(DOCX)
    print(KEY_COMPARE)


if __name__ == "__main__":
    main()
