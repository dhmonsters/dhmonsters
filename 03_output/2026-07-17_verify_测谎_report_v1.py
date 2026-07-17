# 测谎 정적 분석 보고서의 구조와 렌더 결과를 검증하는 스크립트
from pathlib import Path
import ast
import zipfile

from docx import Document
from pypdf import PdfReader


BASE = Path(r"C:\Users\PC\Desktop\02_work\05_AI\03_output")
DOCX_PATH = BASE / "2026-07-17_测谎_파일분석_v1.docx"
PDF_PATH = (
    BASE
    / "2026-07-17_测谎_파일분석_v1_render"
    / "2026-07-17_测谎_파일분석_v1.pdf"
)


document = Document(DOCX_PATH)
all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
all_text += "\n" + "\n".join(
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
)

required_text = [
    "PolygraphBot_vmp.exe",
    "UnrealDbg.aes",
    "UnrealDbgDll.dll",
    "AIHelper.dll",
    "VT_Driver.sys",
    "23.140.4.181",
    "높음",
]
missing_text = [text for text in required_text if text not in all_text]
placeholders = [
    text
    for text in ["TODO", "TBD", "PLACEHOLDER", "{{", "}}"]
    if text in all_text
]

with zipfile.ZipFile(DOCX_PATH) as archive:
    bad_member = archive.testzip()
    document_xml = archive.read("word/document.xml").decode("utf-8")

scripts = [
    BASE / "2026-07-17_测谎_pe_static_analysis_v1.py",
    BASE / "2026-07-17_build_测谎_report_v1.py",
]
for script in scripts:
    ast.parse(script.read_text(encoding="utf-8-sig"))

pdf_pages = len(PdfReader(str(PDF_PATH)).pages)
png_pages = len(list(PDF_PATH.parent.glob("page-*.png")))
cant_split_rows = document_xml.count("<w:cantSplit")
numbered_items = document_xml.count("<w:numPr")

print(f"DOCX_EXISTS={DOCX_PATH.exists()} SIZE={DOCX_PATH.stat().st_size}")
print(
    f"ZIP_CRC_BAD={bad_member} PARAGRAPHS={len(document.paragraphs)} "
    f"TABLES={len(document.tables)}"
)
print(f"MISSING_REQUIRED={missing_text} PLACEHOLDERS={placeholders}")
print(f"CANT_SPLIT_ROWS={cant_split_rows} NUMPR={numbered_items}")
print(f"PDF_PAGES={pdf_pages} PNG_PAGES={png_pages}")
print("PYTHON_AST=PASS")

assert bad_member is None
assert len(document.tables) == 2
assert not missing_text
assert not placeholders
assert pdf_pages == 11
assert png_pages == 11
assert cant_split_rows > 0
print("FINAL_AUDIT=PASS")
