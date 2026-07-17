# Planet Solver 보고서의 번호 목록 재시작과 홀짝 머리글 설정을 검증
import importlib.util
import re
import unittest
from pathlib import Path

from docx import Document


MODULE_PATH = Path(__file__).with_name("2026-07-17_build_planet_solver_report_v1.py")


def load_builder():
    spec = importlib.util.spec_from_file_location("report_builder", MODULE_PATH)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class ReportNumberingTests(unittest.TestCase):
    def test_separate_numbered_sequences_preserve_source_markers(self):
        builder = load_builder()
        doc = Document()
        builder.style_document(doc)
        builder.build_body(doc, ["1. 첫 번째", "2. 두 번째", "", "중간 문단", "", "1. 새 첫 번째"])

        numbered = [p for p in doc.paragraphs if re.match(r"^\d+\. ", p.text)]

        self.assertEqual(len(numbered), 3)
        self.assertEqual([p.text[:3] for p in numbered], ["1. ", "2. ", "1. "])
        self.assertTrue(all(p._p.pPr.numPr is None for p in numbered))

    def test_odd_even_headers_are_not_split(self):
        builder = load_builder()
        doc = Document()
        builder.style_document(doc)

        self.assertFalse(doc.settings.odd_and_even_pages_header_footer)

    def test_bullets_are_explicit_characters(self):
        builder = load_builder()
        doc = Document()
        builder.style_document(doc)
        builder.build_body(doc, ["- 湲癒몃━??"])

        self.assertEqual(doc.paragraphs[-1].text[:2], "• ")
        self.assertIsNone(doc.paragraphs[-1]._p.pPr.numPr)


if __name__ == "__main__":
    unittest.main()
