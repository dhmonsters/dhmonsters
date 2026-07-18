# MapleHunter 사다리 분석 DOCX의 페이지·스타일·표 구조를 검증
import importlib.util
import unittest
from pathlib import Path

from docx.oxml.ns import qn


MODULE_PATH = Path(__file__).with_name(
    "2026-07-18_build_maplehunter_ladder_report_v1.py"
)


def load_builder():
    spec = importlib.util.spec_from_file_location("ladder_report_builder", MODULE_PATH)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class LadderReportDocumentTests(unittest.TestCase):
    def test_document_uses_letter_page_and_one_inch_margins(self):
        doc = load_builder().build_document()
        section = doc.sections[0]

        self.assertAlmostEqual(section.page_width.inches, 8.5, places=2)
        self.assertAlmostEqual(section.page_height.inches, 11.0, places=2)
        self.assertAlmostEqual(section.left_margin.inches, 1.0, places=2)
        self.assertAlmostEqual(section.right_margin.inches, 1.0, places=2)

    def test_tables_have_fixed_9360_dxa_geometry(self):
        doc = load_builder().build_document()

        self.assertGreaterEqual(len(doc.tables), 3)
        for table in doc.tables:
            tbl_pr = table._tbl.tblPr
            width = tbl_pr.find(qn("w:tblW"))
            indent = tbl_pr.find(qn("w:tblInd"))
            grid_width = sum(
                int(col.get(qn("w:w")))
                for col in table._tbl.tblGrid.findall(qn("w:gridCol"))
            )
            self.assertEqual(width.get(qn("w:w")), "9360")
            self.assertEqual(indent.get(qn("w:w")), "120")
            self.assertEqual(grid_width, 9360)

    def test_document_has_running_header_and_footer(self):
        doc = load_builder().build_document()
        section = doc.sections[0]

        self.assertIn("MAPLEHUNTER", section.header.paragraphs[0].text)
        self.assertIn("Page", section.footer.paragraphs[0].text)
        self.assertFalse(doc.settings.odd_and_even_pages_header_footer)
        self.assertFalse(section.different_first_page_header_footer)

    def test_page_field_nodes_are_wrapped_in_a_run(self):
        doc = load_builder().build_document()
        paragraph = doc.sections[0].footer.paragraphs[0]._p

        self.assertIsNone(paragraph.find(qn("w:fldChar")))
        self.assertIsNotNone(paragraph.find(f"{qn('w:r')}/{qn('w:fldChar')}"))

    def test_masthead_uses_short_title_and_clean_target_metadata(self):
        doc = load_builder().build_document()
        text = [paragraph.text for paragraph in doc.paragraphs]

        self.assertIn("MAPLEHUNTER 사다리 처리 분석", text)
        self.assertFalse(any("`MapleHunter_v3.1.17.exe`" in value for value in text))
        self.assertFalse(any("분석일  분석일" in value for value in text))

    def test_table_header_stays_with_first_body_row(self):
        doc = load_builder().build_document()

        for table in doc.tables:
            self.assertTrue(
                all(
                    paragraph.paragraph_format.keep_with_next is True
                    for cell in table.rows[0].cells
                    for paragraph in cell.paragraphs
                )
            )


if __name__ == "__main__":
    unittest.main()
