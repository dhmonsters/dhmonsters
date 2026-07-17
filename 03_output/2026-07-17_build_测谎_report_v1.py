# 测谎 정적 분석 Markdown을 검수 가능한 Word 보고서로 변환하는 빌더
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


sys.stdout.reconfigure(encoding="utf-8")
BASE = Path(r"C:\Users\PC\Desktop\02_work\05_AI\03_output")
SOURCE = BASE / "2026-07-17_测谎_파일분석_v1.md"
OUTPUT = BASE / "2026-07-17_测谎_파일분석_v1.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D1DA"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def create_numbering(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def clean_inline(text):
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def add_text_paragraph(doc, text, num_id=None, bullet=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4 if num_id else 6)
    paragraph.paragraph_format.line_spacing = 1.25
    if num_id is not None:
        apply_numbering(paragraph, num_id)
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=11)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(clean_inline(text))
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    else:
        set_run_font(run, size=13, bold=True, color=BLUE)
    return paragraph


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def table_widths(headers):
    if len(headers) == 3:
        return [2100, 6200, 1060]
    if len(headers) == 2 and "SHA-256" in headers:
        return [2400, 6960]
    if len(headers) == 2:
        return [2500, 6860]
    base = CONTENT_WIDTH_DXA // len(headers)
    widths = [base] * len(headers)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for col_index, value in enumerate(values):
            cell = cells[col_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.08
            run = paragraph.add_run(clean_inline(value))
            size = 8.4 if "SHA-256" in rows[0] and col_index == 1 else 9.3
            set_run_font(run, size=size, bold=row_index == 0, color=INK)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
        if row_index == 0:
            set_repeat_table_header(table.rows[0])
    set_table_geometry(table, table_widths(rows[0]))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def configure_styles(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_masthead(doc):
    header = doc.sections[0].header.paragraphs[0]
    header.text = "STATIC ANALYSIS BRIEF  |  测谎 PACKAGE"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=8.5, bold=True, color=MUTED)
    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("TECHNICAL FINDINGS")
    set_run_font(run, size=10, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("测谎 폴더 정적 분석 보고서")
    set_run_font(run, size=25, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Unreal 게임 디버깅 계층, PolygraphBot 자동화, 원격 통신과 보안 위험")
    set_run_font(run, size=12.5, color=MUTED)
    metadata = [
        ("분석 대상", r"C:\Users\PC\Downloads\Telegram Desktop\测谎"),
        ("분석 방식", "파일 실행 없는 정적 분석"),
        ("분석일", "2026-07-17"),
        ("규모", "파일 20개 · 234,788,386바이트"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}  ")
        set_run_font(label_run, size=10.5, bold=True, color=INK)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=MUTED)
    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(12)
    lead.paragraph_format.space_after = Pt(12)
    lead.paragraph_format.left_indent = Inches(0.15)
    lead.paragraph_format.right_indent = Inches(0.15)
    p_pr = lead._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = lead.add_run("핵심 판단. 생체 신호 기반 거짓말 탐지기가 아니라, Unreal 게임의 안티디버깅을 우회하는 커널·VT 디버거와 보호된 PolygraphBot 자동화 클라이언트의 결합이다.")
    set_run_font(run, size=11, bold=True, color=INK)


def build():
    doc = Document()
    configure_styles(doc)
    add_masthead(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    index = start
    bullet_num = None
    decimal_num = None
    current_list = None
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            current_list = None
            index += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1)
            current_list = None
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2)
            current_list = None
        elif stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            current_list = None
            continue
        elif stripped.startswith("- "):
            if current_list != "bullet":
                bullet_num = create_numbering(doc, bullet=True)
                current_list = "bullet"
            add_text_paragraph(doc, stripped[2:], num_id=bullet_num, bullet=True)
        elif re.match(r"^\d+\.\s", stripped):
            if current_list != "decimal":
                decimal_num = create_numbering(doc, bullet=False)
                current_list = "decimal"
            add_text_paragraph(doc, re.sub(r"^\d+\.\s+", "", stripped), num_id=decimal_num)
        else:
            add_text_paragraph(doc, stripped)
            current_list = None
        index += 1
    props = doc.core_properties
    props.title = "测谎 폴더 정적 분석 보고서"
    props.subject = "UnrealDbg·PolygraphBot 구성과 보안 위험 정적 분석"
    props.author = "OpenAI Codex"
    props.keywords = "UnrealDbg, PolygraphBot, VMProtect, VT, 커널 디버거, 정적 분석"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
