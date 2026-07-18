# MapleHunter 깊바협·난파선·위바협 정적 분석 결과를 DOCX 보고서로 생성하는 도구
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "2026-07-19_maplehunter_canyon_shipwreck_routes_v1.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "E8EEF5"

REPORT_MARKDOWN = r"""
# MapleHunter 깊바협·난파선·위바협 동선 비교 분석

## 1. 결론 요약

확인된 실제 맵 모듈은 `깊바협2`, `깊바협2_2킬`, `난파선의무덤`, `위바협2`, `위바협2_2킬`이다. 사용자 표현인 `깊바혐2_2킬`과 `위바협2_2`는 각각 `깊바협2_2킬`, `위바협2_2킬`에 대응한다. 같은 실행파일 안에 `위바협2테스트`도 있으나 GUI의 정식 선택 목록에는 포함되지 않아 보조 개발판으로 분류했다.

이 다섯 모듈의 핵심은 화면 전체를 이해하는 경로 탐색기가 아니라, 엔진이 갱신하는 미니맵 캐릭터 좌표 `(current_x, current_y)`를 읽고 미리 정한 X·Y 임계값에 따라 키를 누르는 상태 기반 자동화다. 맵 스크립트는 `MapBase`를 상속하고 `run_routine()` 안에서 좌표 확인, 이동, 광역기, 버프, 실패 복구를 반복한다.

`config.json`에는 이 맵들의 세부 동선이나 공격 순서가 저장되지 않는다. 동선·좌표·대기시간·복구 규칙은 각 `maps/*.py` 모듈에 하드코딩되어 있다. 미니맵 영역은 GUI의 맵별 매핑 또는 모듈이 호출하는 `config_manager.set_minimap_region()`으로 적용된다.

가장 큰 차이는 다음과 같다. `깊바협2` 계열은 A/B 두 사냥 구역과 x=85 사다리 왕복을 중심으로 한다. `난파선의무덤`은 수영, 라이딩 점프, 낙하, 연속 텔레포트를 단계식으로 연결한다. `위바협2`는 넓게 층을 순환하며 한 지점당 블리자드 1회를 쓰고, `위바협2_2킬`은 x=63~65 좁은 세로축에서 같은 지점당 블리자드 2회를 쓰는 압축 루틴이다.

## 2. 분석 범위와 근거 수준

- 대상은 `MapleHunter_v3.1.17.exe`의 PyInstaller CArchive에 포함된 맵 소스, 대응 `.pyc`, GUI 모듈 메타데이터, 맵 가이드 이미지다.
- 원본 EXE와 포함 드라이버는 실행하지 않았다. 입력 성공률이나 실제 사냥 성능은 런타임 실험값이 아니라 코드 구조에 대한 정적 분석 결과다.
- **확인된 사실**은 코드에 직접 존재하는 좌표, 조건, 키 입력, 횟수, 타임아웃이다.
- **해석**은 코드의 목적과 실제 맵 동선을 연결한 설명이며, 게임 물리·지연·키 설정에 따라 결과가 달라질 수 있다.

## 3. 공통 작동 구조

1. GUI에서 맵을 선택하면 대응 모듈의 클래스를 불러온다.
2. 좌표 감지부가 미니맵에서 캐릭터 위치를 갱신하고 맵 모듈은 `engine.current_x`, `engine.current_y`를 읽는다.
3. 맵 모듈이 `pyautogui` 형식의 `keyDown`, `keyUp`, `press` 호출로 방향키, 점프, 텔레포트, 공격, 라이딩을 입력한다.
4. 일시정지 시 `MapBase`가 이동·공격 관련 키를 풀고 재개까지 대기한다.
5. 좌표가 목표 범위에 들어오면 다음 단계로 전환하고, 타임아웃·정지·금지 사다리 조건이면 별도 복구 분기를 실행한다.
6. 한 사이클이 끝나면 현재 좌표를 다시 읽어 루틴을 반복한다.

입력 호출부는 `pyautogui` 인터페이스를 사용한다. 실행파일의 다른 공통 모듈에는 입력 백엔드를 Interception 방식으로 바꾸는 몽키패치 설계도 존재하므로, 실제 실행 환경에서는 같은 맵 코드가 기본 입력 또는 패치된 입력 경로를 탈 수 있다. 다만 이번 문서에서는 드라이버를 실행하거나 설치 상태를 확인하지 않았다.

## 4. 전체 비교표

| 모듈 | 주 사냥 축 | 1사이클 광역기 | 주 이동 | 대표 복구 |
|---|---|---|---|---|
| 깊바협2 | A y≤100, B 110<y<145 | A 2회 + B 2회 | 우측 이동·텔포, x85 사다리, 라이딩 좌측 복귀 | 사다리 우측 점프, up 강제 복구, 바닥 구석 좌측 점프 |
| 깊바협2_2킬 | 깊바협2와 동일 | A 3회 + B 3회 | 일반판과 동일 | 일반판과 동일 |
| 난파선의무덤 | 여러 수중 사냥점 | 첫 사이클 최대 6회, 이후 최대 4회 | 좌우 수영, 라이딩 점프, 낙하, 연속 텔포 | 7초 지연 탈출, 우측 텔포 6회 제한, x72 더블점프 |
| 위바협2 | x62~66 중심, x82 하강점 | 3회 | 시간제 점프 상승, 우측 텔포, 반복 하향 텔포 | 5개 금지 사다리 감시, 상단 갇힘 우측 보행 |
| 위바협2_2킬 | x63~65 세로축 | 6회 | Y목표 점프, 아래 점프, 하향 텔포 4회 | 5개 금지 사다리 감시, 9초 범위 이탈 탈출 |

광역기 수는 정상적으로 모든 단계가 실행됐을 때의 코드상 최대치다. `난파선의무덤` 첫 시작 위치가 삼각형 A 밖이면 첫 제네시스는 생략된다.

## 5. 깊바협2 일반판 상세

### 5.1 위치 판정과 기본 순환

| 상태 | 좌표 조건 | 다음 행동 |
|---|---|---|
| A 구역 | y≤100 | A 제네시스 2회 → 2층 이동 |
| B/2층 | 110<y<145 | x85 사다리 동작 → B 제네시스 2회 → A 복귀 |
| 바닥·기타 | 위 두 범위 외 | 라이딩 좌측 점프 복귀 |

1. A에서 직전 A 타이머로부터 20.1초가 지날 때까지 기다린다. 기다리는 중 y>94이면 점프해 기준층을 유지한다.
2. 설정된 공격키를 0.15초씩 두 번 누르고 각 시전 뒤 3.05초 대기한다. 첫 타격 시 A 젠 타이머를 갱신한다.
3. 오른쪽을 누른 채 이동하고 x<43에서 텔레포트 키를 한 번 누른다. y>114이면 2층 진입 성공으로 본다.
4. 현재 x가 85보다 크면 왼쪽, 작으면 오른쪽을 누르면서 `down`도 함께 유지한다. x=85이면서 y>125가 되면 사다리에 걸려 내려가기 시작한 것으로 판정한다.
5. y≥127까지 내려간 뒤 `up`을 누른다. x가 85±2 안에 있고 y=125가 연속 두 번 확인되면 B 위치 확보에 성공한다.
6. B에서 직전 B 타이머로부터 20.1초를 맞춰 제네시스 2회를 시전한다. y≥150이면 추락으로 보고 공격을 중단한다.
7. 라이딩키 기본값 `f`를 누르고 왼쪽을 유지하면서 0.2초 간격으로 점프한다. y≤94가 되면 라이딩을 해제하고 A 복귀 성공으로 처리한다.

### 5.2 사다리 성공률을 높이는 방식

이 루틴은 사다리 x좌표에 단순히 도착하는 것만으로 성공 처리하지 않는다. 이동 방향과 `down`을 동시에 유지해 사다리에 자연스럽게 걸리도록 하고, x=85에서 y가 125보다 커지는 실제 하강 변화를 성공 신호로 사용한다. 그 뒤 아주 짧게 y=127까지 내려갔다가 다시 올라오며 y=125를 두 번 확인한다. 즉, 좌표 정렬, 하강 확인, 재상승 확인의 3단계 검증으로 오인식을 줄인다.

### 5.3 복구 분기

- x=85, y≥131 상태가 메인 루프에서 3초 지속되면 모든 이동키를 풀고 오른쪽 0.05초 후 점프, 0.55초 유지로 사다리를 벗어난다.
- 사다리 함수가 실패한 뒤 x=85±2, y=128~155에 있으면 3초를 추가 관찰한다. 최종 y가 128~140이면 `up`을 최대 3초 눌러 y≤125까지 강제로 복구한다.
- A 복귀 중 y=156~158이 3초 지속되면 왼쪽을 누르고 0.2초 간격 점프를 최대 3.5초 반복한다. x가 시작점보다 2 이상 왼쪽으로 이동하거나 y<156이 되면 조기 종료한다.
- 2층 이동은 5초 타임아웃과 약 1초간 좌표 진전 없음 조건을 둔다. 사다리 접근은 5초, 하강 확인은 3초, 상승 확인은 2.5초 타임아웃을 사용한다.

## 6. 깊바협2_2킬 차이

이 변형은 이동·사다리·복귀·복구 로직이 일반판과 사실상 동일하다. 핵심 차이는 각 A/B 사냥점에서 공격키를 두 번이 아니라 세 번 누른다는 점이다.

| 항목 | 깊바협2 | 깊바협2_2킬 |
|---|---|---|
| 위치당 시전 횟수 | 2회 | 3회 |
| 코드 주석 의미 | 집 파괴 + 본체 파괴 | 집 2방 + 몹 1방 |
| 젠 타이머 갱신 | 1타 직후 | 2타 직후 |
| 시전 간 대기 | 매회 3.05초 | 매회 3.05초 |
| 이동·복구 | 기준 동작 | 동일 |

이름의 `2킬`은 총 시전 횟수가 2회라는 뜻이 아니다. 코드상으로는 집을 두 번 공격한 뒤 몬스터를 한 번 더 공격하는 총 3타 구성이다. 일반판보다 한 지점당 약 3.2초 이상 길어지고, 한 사이클 전체에서는 A와 B 각각 한 타가 추가된다.

## 7. 난파선의무덤 상세

GUI 메타데이터의 미니맵 영역은 `(38,159)~(209,357)`이다. 시작지점 A는 단순 사각형이 아니라 y=159~165에서 아래로 갈수록 넓어지는 삼각형으로 판정한다. y=159일 때 x=88 부근이고, y=165일 때 대략 x=72~103 범위다.

### 7.1 첫 사이클 전용 단계

1. 현재 좌표가 A 삼각형 안이면 제자리 제네시스를 한 번 시전한다.
2. 왼쪽을 누르고 점프를 반복해 x≤73, y≤95로 이동한 뒤 제네시스를 한 번 시전한다.
3. 라이딩에 탑승하고 x=65~69를 짧은 좌우 입력으로 보정하면서 2.9초 동안 점프 상승한다.

첫 사이클 이후에는 1~2단계를 건너뛰고 라이딩 점프 상승 시간을 3.8초로 늘려 공통 구간으로 진입한다.

### 7.2 매 사이클 공통 단계

1. 상승 직후 제자리 제네시스를 한 번 시전한다.
2. 오른쪽으로 걸어 x≥101에 도달하면 제네시스를 한 번 시전한다.
3. y≥110까지 낙하를 기다린다. x<113이면 오른쪽 텔레포트를 반복해 하향 텔레포트 자리를 확보한다.
4. `down`을 0.30초 누른 뒤 텔레포트를 4회 빠르게 입력한다.
5. y≥132 바닥층 도달을 확인하고 제네시스를 한 번 시전한다. 아직 위층이면 하향 텔레포트를 다시 시도한다.
6. 라이딩에 탑승하고 왼쪽을 유지한다. 0.85초마다 점프하며 x≤67까지 수영한 뒤 라이딩을 해제하고 제네시스를 한 번 시전한다.
7. 다음 루프에서는 3.8초 라이딩 상승부터 다시 시작한다.

### 7.3 실패 복구와 리셋

- 왼쪽 B 이동, x101 이동, 낙하, 바닥 확인 단계가 7초 넘게 진전되지 않으면 비상 탈출을 실행한다. 먼저 왼쪽으로 4초 이동하고, 이어 오른쪽으로 이동해 y≥110 낙하 구역에 들어갈 때까지 기다린다.
- 하향 텔레포트 자리를 잡기 위한 오른쪽 텔레포트가 6회에 도달해도 x<113이면 첫 사이클 상태로 리셋한다.
- 마지막 왼쪽 수영 중 x=71~73, y=152~165에 2초 이상 머물면 점프 두 번을 0.35초와 0.3초 간격으로 넣어 벽을 벗어난다.
- 수영·낙하 환경의 특성 때문에 절대 좌표 한 점보다 단계별 X/Y 문턱과 시간 제한을 조합한 것이 핵심이다.

## 8. 위바협2 일반판 상세

모듈 자체가 미니맵 영역 `(84,159)~(217,361)`을 `config_manager`에 설정한다. 시작 x=64, 점프 보정 범위 x=62~66, 우측 하강점 x=82, 시작층 y=104~119를 사용한다.

### 8.1 정상 사이클

1. 보통 x=64로 걸어가 블리자드를 한 번 시전한다. 5번째 사이클마다 x=59를 사용해 같은 자리 반복 패널티를 회피한다.
2. 아래 텔레포트를 한 번 사용한다.
3. 점프키를 2.3초 누르며 x=62~66을 짧은 좌우 탭으로 유지한 뒤 블리자드를 한 번 시전한다.
4. 다시 2.0초 점프 상승하고 0.4초 기다린 뒤 블리자드를 한 번 시전한다.
5. 오른쪽 텔레포트를 한 번 사용하고 걸어서 x=82±2로 보정한다.
6. 최대 15회 동안 매 하향 텔레포트 전에 y와 x를 확인한다. y=104~119에 들어오면 즉시 종료한다.
7. 걸어서 x=64로 복귀하고 다음 사이클을 시작한다.

각 블리자드는 공격키 0.15초 입력 뒤 3.15초를 기다리며, 세 지점 모두 이후 버프 점검을 호출한다.

### 8.2 사다리와 상단 갇힘 복구

백그라운드 감시 스레드는 다음 금지 구간을 1초마다 확인한다.

| 구간 | Y 범위 | 4초 지속 시 탈출 |
|---|---|---|
| x=85~86 | 63~81 | 오른쪽 방향 점프 후 오른쪽 2초 유지 |
| x=77~79 | 86~95 | 오른쪽 방향 점프 후 오른쪽 2초 유지 |
| x=40~42 | 86~95 | 오른쪽 방향 점프 후 오른쪽 2초 유지 |
| x=101~103 | 94~95 | 왼쪽 방향 점프 후 왼쪽 2초 유지 |
| x=98~100 | 93~95 | 왼쪽 방향 점프 후 왼쪽 2초 유지 |

탈출 후 `_ladder_stuck` 플래그를 세워 메인 루틴을 중단하고 x=82 하향 텔레포트 지점으로 이동한 뒤 시작층에 내려가 x=64로 복귀한다. 별도로 x=64~90, y=84~85 상단에 갇히면 오른쪽으로 x≥92까지 걸어 낙하를 유도하고, y>85가 확인되면 x=82로 다시 보정한다.

## 9. 위바협2_2킬 상세

이 변형은 일반판의 넓은 순환을 버리고 x=63~65 세로축을 유지한다. 기준 Y는 시작/2층 114, 1층 136, 상단 93이다. 이름과 달리 한 사이클 전체 광역기는 2회씩 세 구간, 총 6회다.

### 9.1 정상 사이클

1. x=63~65로 보정하고 블리자드를 연속 두 번 시전한 뒤 버프를 점검한다.
2. x를 다시 보정하고 아래 텔레포트 한 번을 사용한다. y≥136을 기다리는 시간은 최대 0.2초다.
3. 점프키를 계속 누르며 x를 짧게 보정하고 y≤93이 될 때까지 상승한다.
4. 상단에서 블리자드를 연속 두 번 시전하고 버프를 점검한다.
5. 0.4초 후 `down+점프`로 아래 점프하고 2초 대기한다. 이어 아래 텔레포트 한 번, 0.5초 대기 후 블리자드를 연속 두 번 시전한다.
6. x를 한 번 보정한 뒤 `down`을 유지하면서 텔레포트를 4회 사용하고 다음 루프로 돌아간다.

### 9.2 복구 방식

- 금지 사다리는 x=71~72 전체 Y, x=77~79/y87~95, x=40~42/y86~95, x=101~103/y94~95, x=98~100/y93~95다. 4초 지속 시 모두 왼쪽 방향 점프 한 번으로 탈출한다.
- x가 63~65 밖이고 Y 변화가 2 이하인 상태가 9초 지속되면 오른쪽 방향 점프를 두 번 수행하고 x를 다시 보정한다.
- 함수 설명문에는 4초라고 적혀 있지만 실제 조건식은 9.0초다. **확인된 사실**은 실행 조건 9초이며, 4초 표기는 오래된 주석으로 해석된다.
- `_adjust_x()`와 y≤93까지의 상승에는 자체 타임아웃이 없다. 좌표 감지가 멈추거나 물리적으로 목표에 도달하지 못하면 백그라운드 사다리 감시 또는 외부 중지에 의존할 가능성이 있다.

## 10. 위바협2 두 버전의 핵심 차이

| 비교 항목 | 위바협2 | 위바협2_2킬 |
|---|---|---|
| X 운용 | x64 중심, x59 회피, x82 하강점 | x63~65 고정 세로축 |
| 상승 종료 | 2.3초와 2.0초 시간제 | y≤93 좌표 도달형 |
| 광역기 | 세 위치에서 각 1회, 총 3회 | 세 위치에서 각 2회, 총 6회 |
| 하강 | x82에서 시작층 y104~119까지 최대 15회 | 아래 점프·텔포 후 마지막에 텔포 4회 |
| 사다리 탈출 | 방향 점프 + 같은 방향 2초 유지, 플래그 복귀 | 왼쪽 방향 점프 1회, 별도 플래그 없음 |
| 상단 갇힘 | x64~90/y84~85 전용 복구 있음 | 전용 복구 없음 |
| 범위 이탈 | 전용 9초 검사 없음 | x63~65 이탈·Y정지 9초 검사 |
| 제자리 회피 | 5사이클마다 x59 | 없음 |

일반판은 각 층을 넓게 훑고 복구를 강하게 건 안정형이다. 2킬판은 좁은 X축에서 공격 횟수를 늘리고 이동을 줄인 화력 집중형이다. 따라서 2킬판은 정확한 좌표 감지와 점프 상승 성공에 더 민감하고, 일반판은 이동량이 많은 대신 실패 지점별 복구가 더 촘촘하다.

## 11. config.json과 맵 선택의 관계

`config.json`에서 기대할 수 있는 것은 미니맵 영역, 사냥·감지 설정 같은 공통 값이다. 이 다섯 맵의 순서, x=85 사다리, x=101 사냥점, 블리자드 횟수, 텔레포트 횟수는 `config.json`이 아니라 모듈 소스에 직접 들어 있다.

GUI 메타데이터에는 `깊바협2`의 미니맵 영역 `(58,161)~(220,362)`, `난파선의무덤`의 `(38,159)~(209,357)`, `위바협2`의 `(84,159)~(217,361)`가 연결되어 있다. `위바협2` 두 모듈은 시작할 때 같은 영역을 다시 설정한다. 따라서 설정 파일만 비교해서는 버전별 루틴 차이를 찾을 수 없고, 실제 차이는 각 맵 클래스의 `run_routine()`과 보조 함수에서 확인해야 한다.

## 12. 안정성과 유지보수 관점

- 좌표는 미니맵 픽셀 좌표에 가깝기 때문에 창 크기, 미니맵 위치, 감지 오차가 바뀌면 임계값 기반 분기가 틀어질 수 있다.
- `keyDown` 뒤 `finally`에서 `keyUp`을 보장하는 함수가 많아 이동키 고착을 줄인다. 일시정지 때 전체 키를 두 번 해제하는 `MapBase`도 같은 목적이다.
- 일부 보정 루프에는 타임아웃이 없고, 백그라운드 사다리 감시 스레드가 메인 루틴과 동시에 방향키를 조작한다. 이론상 두 경로가 같은 순간 키를 누르거나 풀 수 있다.
- 주석과 실제 코드가 다른 곳이 있다. 대표적으로 `위바협2_2킬` 범위 이탈은 설명 4초, 실제 9초이며, 루틴 주석의 0.2초 대기와 실제 `time.sleep(0.4)`도 다르다. 유지보수 시 조건식과 호출값을 우선해야 한다.
- `깊바협2_2킬`은 이동 최적화판이 아니라 공격 1회를 각 사냥점에 추가한 변형이다. 반면 `위바협2_2킬`은 이동 구조 자체를 세로축 중심으로 크게 바꿨다.

## 13. 최종 판단

세 맵 계열 모두 핵심은 미니맵 좌표 피드백과 시간제 키 입력의 결합이다. 가장 정교한 단일 동작은 `깊바협2`의 x=85 사다리 탑승 검증이며, 가장 긴 상태 흐름은 `난파선의무덤`의 첫 사이클/반복 사이클 분리다. 가장 큰 버전 차이는 `위바협2`와 `위바협2_2킬` 사이에 있다.

실제 성공률을 높이려면 먼저 미니맵 영역과 `(x,y)` 감지가 안정적인지 확인해야 한다. 그다음 사다리·상단·수중 벽처럼 코드가 별도 복구 조건을 둔 좌표에서 로그를 대조하는 것이 효율적이다. 공격 횟수보다 좌표 오검출과 키 해제 경합이 전체 루틴 실패에 더 직접적인 영향을 줄 수 있다.
"""


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("MAPLEHUNTER | CANYON & SHIPWRECK ROUTE ANALYSIS"), size=9, color=GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_run = footer.add_run("Page ")
    set_run_font(page_run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    page_run._r.append(begin)
    page_run._r.append(instruction)
    page_run._r.append(end)


def add_inline(paragraph, text, size=11):
    for part in re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=max(8.5, size - 1))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table_widths(column_count):
    presets = {
        3: [1800, 2200, 5360],
        4: [1700, 2600, 1900, 3160],
        5: [1500, 1700, 1700, 1900, 2560],
    }
    if column_count in presets:
        return presets[column_count]
    base = 9360 // column_count
    widths = [base] * column_count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc, rows):
    header = rows[0]
    body = rows[2:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    font_size = 8.4 if len(header) >= 5 else 9.3
    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        shade_cell(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        add_inline(paragraph, value, size=font_size)
        for run in paragraph.runs:
            run.bold = True
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)

    for values in body:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            add_inline(paragraph, value, size=font_size)
    apply_table_geometry(table, table_widths(len(header)))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_masthead(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("MAPLEHUNTER 수중 맵 동선 비교"), size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("깊바협2·난파선의무덤·위바협2 | 일반판과 2킬판 정적 분석"), size=13, color=GRAY)
    for label, value in (
        ("분석일", "2026-07-19"),
        ("대상", "MapleHunter_v3.1.17.exe"),
        ("방식", "실행하지 않은 PyInstaller 맵 소스·GUI 메타데이터 정적 분석"),
        ("결론", "미니맵 좌표 피드백 + 시간제 키 입력 + 맵별 복구 분기"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}  "), size=10.5, bold=True)
        set_run_font(paragraph.add_run(value), size=10.5, color=GRAY)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(4)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    borders.append(bottom)
    rule._p.get_or_add_pPr().append(borders)


def create_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    base_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = numbering.xpath(f'./w:num[@w:numId="{base_num_id}"]')[0]
    abstract_id = base_num.xpath("./w:abstractNumId")[0].get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def set_paragraph_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(number)
    p_pr.append(num_pr)


def add_body(doc, lines):
    index = 1
    active_numbering = None
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            active_numbering = None
            index += 1
            continue
        if line.startswith("## "):
            active_numbering = None
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            active_numbering = None
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("|"):
            active_numbering = None
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(doc, rows)
            continue
        if re.match(r"^\d+\.\s", line):
            if active_numbering is None:
                active_numbering = create_numbering_instance(doc)
            paragraph = doc.add_paragraph(style="List Number")
            set_paragraph_numbering(paragraph, active_numbering)
            add_inline(paragraph, re.sub(r"^\d+\.\s*", "", line))
            index += 1
            continue
        if line.startswith("- "):
            active_numbering = None
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, line[2:].strip())
            index += 1
            continue

        active_numbering = None
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith(("## ", "### ", "|", "- ")) or re.match(r"^\d+\.\s", next_line):
                break
            paragraph_lines.append(next_line)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def build_report(output_path=OUTPUT):
    output_path = Path(output_path)
    doc = Document()
    configure_document(doc)
    add_masthead(doc)
    add_body(doc, REPORT_MARKDOWN.strip().splitlines())
    doc.core_properties.title = "MapleHunter 깊바협·난파선·위바협 동선 비교 분석"
    doc.core_properties.subject = "MapleHunter v3.1.17 underwater map route static analysis"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.keywords = "MapleHunter, 깊바협2, 난파선의무덤, 위바협2, static analysis"
    doc.save(output_path)


if __name__ == "__main__":
    build_report()
    print(OUTPUT)
