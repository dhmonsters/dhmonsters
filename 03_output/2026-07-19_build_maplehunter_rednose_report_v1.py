# 빨코2 전 버전 비교 분석 Markdown을 검증 가능한 DOCX로 변환하는 도구
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "2026-07-19_maplehunter_rednose_route_v1.md"
OUTPUT = ROOT / "2026-07-19_maplehunter_rednose_route_v1.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "E8EEF5"


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("MAPLEHUNTER | REDCO2 ROUTE ANALYSIS"), size=9, color=GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    page_run = footer.add_run("Page ")
    set_run_font(page_run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    page_run._r.append(begin)
    page_run._r.append(instruction)
    page_run._r.append(end)


def add_inline(paragraph, text, size=11):
    for part in re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=max(8.5, size - 1))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table_widths(column_count):
    if column_count == 4:
        return [1700, 2600, 1900, 3160]
    if column_count == 7:
        return [1120, 1300, 1240, 1240, 1240, 1240, 1980]
    base = 9360 // column_count
    widths = [base] * column_count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc, rows):
    header = rows[0]
    body = rows[2:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    font_size = 8.2 if len(header) >= 6 else 9.5
    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        shade_cell(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        add_inline(paragraph, value, size=font_size)
        for run in paragraph.runs:
            run.bold = True
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)

    for values in body:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            add_inline(paragraph, value, size=font_size)
    apply_table_geometry(table, table_widths(len(header)))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_masthead(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("MAPLEHUNTER 빨코2 전 버전 비교"), size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("v3.1.17 | 일반형·플래닛 v2~v5·빨코new 동선과 복구 방식"), size=13, color=GRAY)
    for label, value in (
        ("분석일", "2026-07-19"),
        ("대상", "MapleHunter_v3.1.17.exe"),
        ("방식", "실행하지 않은 PyInstaller 맵 소스·코드 객체 정적 분석"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}  "), size=10.5, bold=True)
        set_run_font(paragraph.add_run(value), size=10.5, color=GRAY)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(4)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    borders.append(bottom)
    rule._p.get_or_add_pPr().append(borders)


def create_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    base_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = numbering.xpath(f'./w:num[@w:numId="{base_num_id}"]')[0]
    abstract_id = base_num.xpath("./w:abstractNumId")[0].get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def set_paragraph_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(number)


def add_body(doc, lines):
    index = 1
    active_numbering = None
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            active_numbering = None
            index += 1
            continue
        if line.startswith("## "):
            active_numbering = None
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            active_numbering = None
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("|"):
            active_numbering = None
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(doc, rows)
            continue
        if re.match(r"^\d+\.\s", line):
            if active_numbering is None:
                active_numbering = create_numbering_instance(doc)
            paragraph = doc.add_paragraph(style="List Number")
            set_paragraph_numbering(paragraph, active_numbering)
            add_inline(paragraph, re.sub(r"^\d+\.\s*", "", line))
            index += 1
            continue
        if line.startswith("- "):
            active_numbering = None
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, line[2:].strip())
            index += 1
            continue

        active_numbering = None
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith(("## ", "### ", "|", "- ")) or re.match(r"^\d+\.\s", next_line):
                break
            paragraph_lines.append(next_line)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def build_report(source_path, output_path):
    source_path = Path(source_path)
    output_path = Path(output_path)
    lines = source_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_masthead(doc)
    add_body(doc, lines)
    doc.core_properties.title = "MapleHunter 빨코2 전 버전 동선 비교 분석"
    doc.core_properties.subject = "MapleHunter v3.1.17 redco2 route static analysis"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.keywords = "MapleHunter, 빨코2, route, static analysis"
    doc.save(output_path)


if __name__ == "__main__":
    build_report(SOURCE, OUTPUT)
    print(OUTPUT)
