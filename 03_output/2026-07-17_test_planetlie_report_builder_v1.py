# Planetlie Word 보고서 생성기의 OOXML 유효성을 검사하는 회귀 테스트
import importlib.util
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


BUILDER = Path(__file__).with_name("2026-07-17_build_planetlie_report_v1.py")
spec = importlib.util.spec_from_file_location("planetlie_report_builder", BUILDER)
module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(module)


def test_page_field_nodes_are_wrapped_in_runs():
    document = Document()
    paragraph = document.add_paragraph()
    module.add_page_field(paragraph)
    field_nodes = paragraph._p.xpath(".//w:fldChar | .//w:instrText")
    assert field_nodes, "PAGE 필드 노드가 생성되지 않았습니다."
    for node in field_nodes:
        assert node.getparent().tag == qn("w:r"), (
            f"{node.tag} 노드는 Word 규칙에 따라 w:r 안에 있어야 합니다."
        )


def test_lists_use_bullets_and_restart_numbering_by_section():
    output = module.build_report()
    document = Document(output)
    paragraphs = {paragraph.text.strip(): paragraph for paragraph in document.paragraphs}

    bullet = next(
        paragraph
        for text, paragraph in paragraphs.items()
        if text.startswith("Microsoft Visual C++")
    )
    assert bullet.style.name == "List Bullet", "일반 항목은 실제 글머리표 스타일이어야 합니다."

    numbered_starts = [
        next(
            paragraph
            for text, paragraph in paragraphs.items()
            if text.startswith(prefix)
        )
        for prefix in (
            "nika.exe가 사용자 화면",
            "게임의 GameAssembly.dll",
            "nika.exe. 전체 실행 흐름",
        )
    ]
    num_ids = [paragraph._p.pPr.numPr.numId.val for paragraph in numbered_starts]
    assert len(set(num_ids)) == len(num_ids), "서로 다른 번호 목록은 1부터 다시 시작해야 합니다."


def test_numbering_definitions_follow_wordprocessingml_order():
    output = module.build_report()
    document = Document(output)
    numbering = document.part.numbering_part.element
    child_tags = [child.tag for child in numbering]
    abstract_indices = [
        index for index, tag in enumerate(child_tags) if tag == qn("w:abstractNum")
    ]
    num_indices = [index for index, tag in enumerate(child_tags) if tag == qn("w:num")]
    assert max(abstract_indices) < min(num_indices), (
        "모든 abstractNum 정의는 WordprocessingML 규칙상 num 인스턴스보다 앞에 있어야 합니다."
    )


if __name__ == "__main__":
    test_page_field_nodes_are_wrapped_in_runs()
    test_lists_use_bullets_and_restart_numbering_by_section()
    test_numbering_definitions_follow_wordprocessingml_order()
    print("PASS")
