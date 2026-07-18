# 빨코2 전 버전 비교 DOCX 빌더의 내용과 문서 구조를 검증하는 테스트
import importlib.util
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


MODULE_PATH = Path(__file__).with_name("2026-07-19_build_maplehunter_rednose_report_v1.py")
SOURCE_PATH = Path(__file__).with_name("2026-07-19_maplehunter_rednose_route_v1.md")


def load_target():
    spec = importlib.util.spec_from_file_location("rednose_report_builder", MODULE_PATH)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class RednoseReportBuilderTest(unittest.TestCase):
    def build_document(self):
        module = load_target()
        tmp = tempfile.TemporaryDirectory()
        output = Path(tmp.name) / "report.docx"
        module.build_report(SOURCE_PATH, output)
        return tmp, output

    def test_contains_every_requested_version_and_route_summary(self):
        tmp, output = self.build_document()
        try:
            text = "\n".join(p.text for p in Document(output).paragraphs)
            for label in ["일반 빨코2", "빨코2(플래닛)", "빨코2_v2", "빨코2_v3", "빨코2_v4", "빨코2_v5", "빨코new"]:
                self.assertIn(label, text)
            self.assertIn("최종 동선 요약", text)
        finally:
            tmp.cleanup()

    def test_uses_letter_page_and_one_inch_margins(self):
        tmp, output = self.build_document()
        try:
            section = Document(output).sections[0]
            self.assertEqual(round(section.page_width.inches, 2), 8.5)
            self.assertEqual(round(section.page_height.inches, 2), 11.0)
            for margin in [section.top_margin, section.right_margin, section.bottom_margin, section.left_margin]:
                self.assertEqual(round(margin.inches, 2), 1.0)
        finally:
            tmp.cleanup()

    def test_tables_have_fixed_full_width_geometry(self):
        tmp, output = self.build_document()
        try:
            doc = Document(output)
            self.assertGreaterEqual(len(doc.tables), 2)
            for table in doc.tables:
                tbl_pr = table._tbl.tblPr
                self.assertEqual(tbl_pr.find(qn("w:tblW")).get(qn("w:w")), "9360")
                self.assertEqual(tbl_pr.find(qn("w:tblInd")).get(qn("w:w")), "120")
                self.assertEqual(sum(int(col.get(qn("w:w"))) for col in table._tbl.tblGrid.gridCol_lst), 9360)
        finally:
            tmp.cleanup()

    def test_footer_page_field_is_inside_a_run(self):
        tmp, output = self.build_document()
        try:
            with zipfile.ZipFile(output) as archive:
                footer = archive.read("word/footer1.xml").decode("utf-8")
            self.assertIn("PAGE", footer)
            self.assertNotIn("<w:p><w:fldChar", footer)
            run_start = footer.index("<w:r>")
            run_end = footer.index("</w:r>", run_start)
            field_at = footer.index("<w:fldChar", run_start)
            self.assertLess(field_at, run_end)
        finally:
            tmp.cleanup()

    def test_each_numbered_sequence_uses_a_fresh_numbering_instance(self):
        tmp, output = self.build_document()
        try:
            doc = Document(output)
            starts = [
                "일반 빨코2와 레거시",
                "방향키 누름.",
                "2층 -2.9",
                "config.json의 CCTV",
            ]
            num_ids = []
            for start in starts:
                paragraph = next(p for p in doc.paragraphs if p.text.startswith(start))
                num_pr = paragraph._p.pPr.numPr
                if num_pr is None:
                    num_pr = paragraph.style._element.pPr.numPr
                num_ids.append(num_pr.numId.val)
            self.assertEqual(len(set(num_ids)), len(num_ids))

            numbering = doc.part.numbering_part.element
            for num_id in num_ids:
                num = numbering.find(f".//{{{qn('w:num').split('}')[0][1:]}}}num[@{qn('w:numId')}='{num_id}']")
                self.assertIsNotNone(num)
                start_override = num.find(f".//{qn('w:startOverride')}")
                self.assertIsNotNone(start_override)
                self.assertEqual(start_override.get(qn("w:val")), "1")
        finally:
            tmp.cleanup()


if __name__ == "__main__":
    unittest.main()
