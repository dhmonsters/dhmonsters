# Planetlie 정적 분석 결과를 전문적인 Word 보고서로 만드는 도구
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(r"C:\Users\PC\Desktop\02_work\05_AI\03_output")
SOURCE_MD = OUTPUT_DIR / "2026-07-17_planetlie_파일분석_v1.md"
OUTPUT_DOCX = OUTPUT_DIR / "2026-07-17_planetlie_파일분석_v1.docx"
SAMPLE_IMAGE = Path(
    r"C:\Users\PC\Downloads\planetliev1.02\planetlie\PIC\demo22_det_frame_0.jpg"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "16324F"
MUTED = "667085"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "EEF4FA"
WHITE = "FFFFFF"
BLACK = "111111"


def set_font(run, size=None, bold=None, color=None, italic=None, mono=False):
    latin_font = "Consolas" if mono else "Calibri"
    east_asia_font = "Malgun Gothic"
    run.font.name = latin_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet_style = doc.styles["List Bullet"]
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(11)
    bullet_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    bullet_style.paragraph_format.left_indent = Inches(0.5)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
    bullet_style.paragraph_format.space_after = Pt(8)
    bullet_style.paragraph_format.line_spacing = 1.167


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for child in (begin, instr, separate, text, end):
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("PLANETLIE V1.02  |  STATIC ANALYSIS")
    set_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    add_page_field(p)
    for run in p.runs:
        set_font(run, size=8.5, color=MUTED)


def next_numbering_id(numbering, tag):
    values = [
        int(element.get(qn(tag)))
        for element in numbering
        if element.get(qn(tag)) is not None
    ]
    return max(values, default=0) + 1


def make_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNumId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    first_num_index = next(
        (
            index
            for index, child in enumerate(numbering)
            if child.tag == qn("w:num")
        ),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num_id = next_numbering_id(numbering, "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])


def add_inline(paragraph, text, size=11, color=BLACK, bold=False):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size - 0.5, color="3B4652", mono=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, color=color, bold=bold)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    cell_margins = tbl_pr.find(qn("w:tblCellMar"))
    if cell_margins is None:
        cell_margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = cell_margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            cell_margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, rows):
    header = rows[0]
    if len(header) == 3:
        widths = [1728, 3312, 4320]
        font_size = 9.2
    else:
        widths = [1440, 2016, 864, 5040]
        font_size = 8.3
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, value, size=font_size, bold=True)
    for row_data in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(header) == 4 and index == 2
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            add_inline(p, value, size=font_size)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)


def add_callout(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.line_spacing = 1.10
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shading)
    run = p.add_run("핵심 요약  ")
    set_font(run, size=11, bold=True, color=NAVY)
    run = p.add_run(
        "nika.exe가 지휘하고 HDDebug.dll이 자동화 기능을 제공하며, "
        "mxdin.dll이 게임 내부 Lua 실행 통로를 만드는 3단 구조입니다."
    )
    set_font(run, size=11, color=BLACK)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("정적 분석 보고서")
    set_font(run, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Planetlie v1.02")
    set_font(run, size=27, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("사용 도구, 핵심 구성요소와 전체 작동 방식")
    set_font(run, size=14, color=MUTED)

    metadata = [
        ("분석 대상", r"C:\Users\PC\Downloads\planetliev1.02\planetlie"),
        ("분석일", "2026-07-17"),
        ("방식", "파일을 실행하지 않은 정적 분석"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{label}  ")
        set_font(run, size=10.5, bold=True, color=DARK_BLUE)
        add_inline(p, value, size=10.5, color=BLACK)
    add_callout(doc)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].startswith("|"):
        parts = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            rows.append(parts)
        index += 1
    return rows, index


def build_report():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    index = 5
    image_inserted = False
    active_list_kind = None
    number_id = None
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("## "):
            active_list_kind = None
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:], size=16, color=BLUE, bold=True)
        elif line.startswith("### "):
            active_list_kind = None
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], size=13, color=BLUE, bold=True)
        elif line.startswith("| "):
            active_list_kind = None
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        elif re.match(r"^\d+\. ", line):
            if active_list_kind != "number":
                number_id = make_numbering(doc, "decimal")
            active_list_kind = "number"
            p = doc.add_paragraph()
            apply_numbering(p, number_id)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            active_list_kind = "bullet"
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        else:
            active_list_kind = None
            p = doc.add_paragraph()
            add_inline(p, line)
            if (
                not image_inserted
                and "샘플 이미지의 “투명 도형 찾기 준비”" in line
                and SAMPLE_IMAGE.exists()
            ):
                image_p = doc.add_paragraph()
                image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_p.paragraph_format.space_before = Pt(6)
                image_p.add_run().add_picture(str(SAMPLE_IMAGE), width=Inches(6.25))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                run = caption.add_run(
                    "그림 1. 패키지에 포함된 MapleStory Worlds 투명 도형 탐지 준비 화면"
                )
                set_font(run, size=9, italic=True, color=MUTED)
                image_inserted = True
        index += 1

    doc.core_properties.title = "Planetlie v1.02 정적 분석 보고서"
    doc.core_properties.subject = "사용 프로그램, 핵심 구성요소와 작동 방식"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.keywords = "Planetlie, static analysis, OpenCV, HDDebug, IL2CPP"
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_report())
