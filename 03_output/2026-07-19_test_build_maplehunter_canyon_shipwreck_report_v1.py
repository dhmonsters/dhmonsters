# MapleHunter 깊바협·난파선·위바협 비교 보고서 생성 결과를 검증하는 테스트
import importlib.util
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
BUILDER = ROOT / "2026-07-19_build_maplehunter_canyon_shipwreck_report_v1.py"


class CanyonShipwreckReportBuilderTest(unittest.TestCase):
    def _load_builder(self):
        self.assertTrue(BUILDER.exists(), f"보고서 생성기가 아직 없습니다. {BUILDER.name}")
        spec = importlib.util.spec_from_file_location("canyon_shipwreck_builder", BUILDER)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return module

    def test_builds_report_with_all_requested_maps_and_core_findings(self):
        module = self._load_builder()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "report.docx"
            module.build_report(output)
            self.assertTrue(output.exists())

            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)
            for required in (
                "깊바협2",
                "깊바협2_2킬",
                "난파선의무덤",
                "위바협2",
                "위바협2_2킬",
                "사다리",
                "텔레포트",
                "config.json",
                "확인된 사실",
                "해석",
            ):
                self.assertIn(required, text)

    def test_uses_letter_geometry_real_numbering_and_fixed_tables(self):
        module = self._load_builder()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "report.docx"
            module.build_report(output)
            doc = Document(output)
            section = doc.sections[0]
            self.assertEqual(round(section.page_width.inches, 2), 8.50)
            self.assertEqual(round(section.page_height.inches, 2), 11.00)
            self.assertEqual(round(section.left_margin.inches, 2), 1.00)
            self.assertEqual(round(section.right_margin.inches, 2), 1.00)

            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
                footer_xml = "".join(
                    archive.read(name).decode("utf-8")
                    for name in archive.namelist()
                    if name.startswith("word/footer") and name.endswith(".xml")
                )
            self.assertRegex(document_xml, r'<w:tblW[^>]*(?:w:w="9360"[^>]*w:type="dxa"|w:type="dxa"[^>]*w:w="9360")')
            self.assertRegex(document_xml, r'<w:tblInd[^>]*(?:w:w="120"[^>]*w:type="dxa"|w:type="dxa"[^>]*w:w="120")')
            self.assertIn("w:numPr", document_xml)
            self.assertIn("w:abstractNum", numbering_xml)
            self.assertIn("PAGE", footer_xml)

    def test_contains_no_placeholders_or_internal_citation_tokens(self):
        module = self._load_builder()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "report.docx"
            module.build_report(output)
            with zipfile.ZipFile(output) as archive:
                xml = archive.read("word/document.xml").decode("utf-8")
            for forbidden in ("TODO", "PLACEHOLDER", "{{", "[[", "turn0", "citation"):
                self.assertNotIn(forbidden, xml)

    def test_shipwreck_repeated_cycle_counts_four_genesis_casts(self):
        module = self._load_builder()
        self.assertIn("이후 최대 4회", module.REPORT_MARKDOWN)
        self.assertNotIn("이후 최대 5회", module.REPORT_MARKDOWN)


if __name__ == "__main__":
    unittest.main()
