# Planet Solver 정적 분석 Markdown 원고를 전문적인 Word 보고서로 변환하는 도구
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "2026-07-17_planet_solver_static_analysis_v1.md"
OUTPUT = ROOT / "2026-07-17_planet_solver_static_analysis_v1.docx"
IMAGE_ROOT = ROOT / "2026-07-17_planet_solver_static_analysis_v1_extracted" / "MapleHunter_v3.1.17.exe" / "detectors" / "planet"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "F2F4F7"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_manual_list_item(paragraph):
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_inline(paragraph, text, size=11, color=None):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=max(8.5, size - 1), color=RGBColor(0x33, 0x33, 0x33))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("PLANET SOLVER | STATIC ANALYSIS")
    set_run_font(run, size=8.5, color=GRAY, bold=True)
    add_page_field(section.footer.paragraphs[0])


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("PLANET SOLVER 정적 분석")
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("v1.0.5 및 MapleHunter v3.1.17의 구조·작동 방식·보안 위험")
    set_run_font(run, size=13.5, color=GRAY)

    metadata = [
        ("분석일", "2026-07-17"),
        ("대상", r"C:\Users\PC\Downloads\Telegram Desktop\sssa\플래닛 (2)"),
        ("방식", "실행하지 않은 정적 분석"),
        ("상태", "기술 구조 확인 완료 · 서버 동적 코드 영역은 미확인"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}  ")
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF2F8")
    p_pr.append(shd)
    run = p.add_run("한줄 판단  로컬 YOLO 추론과 원격 동적 코드를 결합한 게임 자동화 도구이며, 미서명 파일·보안 비활성화 요구·커널 입력 드라이버 때문에 주 PC 실행 위험이 높습니다.")
    set_run_font(run, size=11, color=DARK_BLUE, bold=True)


def table_widths(headers):
    key = tuple(headers)
    if key == ("파일", "크기", "서명", "판단된 역할"):
        return [3100, 1500, 1350, 3410]
    if key == ("구성요소", "사용 주체", "역할"):
        return [2350, 1650, 5360]
    if key == ("위험", "수준", "근거와 의미"):
        return [2350, 1100, 5910]
    count = len(headers)
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    headers = [cell.strip() for cell in rows[0]]
    data_rows = rows[2:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = table_widths(headers)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, header, size=9.5)
        for run in p.runs:
            run.bold = True

    for row_values in data_rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value.strip(), size=9.2)
            if headers[index] in {"크기", "서명", "수준", "사용 주체"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_figure(doc, path, caption, width):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    run = cp.add_run(caption)
    set_run_font(run, size=9, color=GRAY, italic=True)


def build_body(doc, lines):
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            doc.add_paragraph(heading, style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            heading = line[4:].strip()
            doc.add_paragraph(heading, style="Heading 2")
            if heading.endswith("투명 도형 트리거"):
                add_figure(doc, IMAGE_ROOT / "debug_puzzle.png", "그림 1. 투명 도형 추적 퍼즐의 디버그 캡처", 5.7)
            elif heading.endswith("일반 거짓말 탐지기 풀이"):
                add_figure(doc, IMAGE_ROOT / "debug_lie.png", "그림 2. OCR 방식으로 처리하는 4자리 거짓말 탐지기", 4.8)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and set(lines[index + 1].replace("|", "").strip()) <= {"-", ":", " "}:
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if re.match(r"^- ", line):
            p = doc.add_paragraph()
            style_manual_list_item(p)
            p.add_run("• ")
            add_inline(p, line[2:].strip())
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            style_manual_list_item(p)
            marker = re.match(r"^(\d+\.) ", line).group(1)
            p.add_run(f"{marker} ")
            add_inline(p, re.sub(r"^\d+\. ", "", line).strip())
            index += 1
            continue

        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines):
            next_line = lines[index].rstrip()
            if not next_line or next_line.startswith(("#", "|", "- ")) or re.match(r"^\d+\. ", next_line):
                break
            paragraph_lines.append(next_line.strip())
            index += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))


def main():
    doc = Document()
    style_document(doc)
    add_title_block(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    first_section = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    build_body(doc, lines[first_section:])
    doc.core_properties.title = "Planet Solver v1.0.5 정적 분석 보고서"
    doc.core_properties.subject = "작동 방식, 사용 기술, 프로그램 관계와 보안 위험"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.keywords = "Planet Solver, MapleHunter, PyInstaller, ncnn, YOLO, 정적 분석"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
